"""
RAG (Retrieval Augmented Generation) Service
Handles document processing, embedding generation, and semantic search
Now integrated with centralized model service
"""

import os
import logging
from pathlib import Path
from typing import List, Dict, Optional, Tuple

try:
    import chromadb
    from chromadb.config import Settings
except Exception as e:
    chromadb = None
    Settings = None
    logger = logging.getLogger(__name__)
    logger.warning("ChromaDB import failed: %s", e)

try:
    from sentence_transformers import SentenceTransformer
except Exception as e:
    SentenceTransformer = None
    logger = logging.getLogger(__name__)
    logger.warning("sentence-transformers import failed: %s", e)

from PyPDF2 import PdfReader
import docx
import pandas as pd
from io import BytesIO
try:
    import tiktoken
    HAS_TIKTOKEN = True
except ImportError:
    tiktoken = None
    HAS_TIKTOKEN = False
import hashlib

logger = logging.getLogger(__name__)

class RAGService:
    """Manages RAG operations: indexing, searching, and retrieval"""
    
    def __init__(self, chroma_path: str = None, model_name: str = None):
        """Initialize RAG service with ChromaDB and embedding model"""
        # Load configuration from FAIA config manager
        try:
            import sys
            from pathlib import Path
            config_path = Path(__file__).parent.parent / "configuration"
            sys.path.append(str(config_path))
            from faia_config_manager import config_manager
            
            rag_config = config_manager.get_rag_config()
            
            # Use config values with fallbacks
            config_path = rag_config.get("vector_db", {}).get("path", "./chroma_db")
            self.chroma_path = Path(chroma_path or config_path)
            self.model_name = model_name or rag_config.get("embedding", {}).get("model", "all-MiniLM-L6-v2")
            self.collection_name = rag_config.get("vector_db", {}).get("collection_name", "course_materials")
            self.chunk_size = rag_config.get("chunking", {}).get("chunk_size", 500)
            self.chunk_overlap = rag_config.get("chunking", {}).get("chunk_overlap", 50)
            self.default_top_k = rag_config.get("retrieval", {}).get("top_k", 5)
            self.max_context_tokens = rag_config.get("retrieval", {}).get("max_context_tokens", 2000)
            self.min_similarity = rag_config.get("retrieval", {}).get("min_similarity", 0.3)
            
            logger.info("RAG service loaded configuration from FAIA config")
        except Exception as e:
            # Fallback to hardcoded values if config fails
            logger.warning("Failed to load RAG config, using defaults: %s", e)
            self.chroma_path = Path(chroma_path or "./chroma_db")
            self.model_name = model_name or "all-MiniLM-L6-v2"
            self.collection_name = "course_materials"
            self.chunk_size = 500
            self.chunk_overlap = 50
            self.default_top_k = 5
            self.max_context_tokens = 2000
            self.min_similarity = 0.3
        
        # Create chroma directory
        self.chroma_path.mkdir(parents=True, exist_ok=True)
        
        # Initialize ChromaDB if available
        if chromadb is None or Settings is None:
            self.chroma_client = None
            self.collection = None
            self.embedding_model = None
            logger.warning("ChromaDB unavailable; using stub RAG behavior")
            return
        
        try:
            self.chroma_client = chromadb.PersistentClient(
                path=str(self.chroma_path),
                settings=Settings(anonymized_telemetry=False)
            )
            
            self.collection = self.chroma_client.get_or_create_collection(
                name=self.collection_name,
                metadata={"description": "FAIA course materials for RAG"}
            )
        except Exception as e:
            self.chroma_client = None
            self.collection = None
            logger.warning("ChromaDB initialization failed: %s", e)

        # Initialize embedding model if available
        if SentenceTransformer is None:
            self.embedding_model = None
            logger.warning("sentence-transformers unavailable; using stub RAG behavior")
            return

        try:
            logger.info("Loading embedding model: %s", self.model_name)
            self.embedding_model = SentenceTransformer(self.model_name)
            logger.info("RAG service initialized")
        except Exception as e:
            self.embedding_model = None
            logger.warning("Embedding model initialization failed: %s", e)
        
        # Initialize tokenizer for counting tokens
        if HAS_TIKTOKEN:
            try:
                self.tokenizer = tiktoken.get_encoding("cl100k_base")
            except Exception as e:
                self.tokenizer = None
                logger.debug("Tokenizer initialization failed: %s", e)
        else:
            self.tokenizer = None
    
    def count_tokens(self, text: str) -> int:
        """Count tokens in text"""
        if self.tokenizer:
            return len(self.tokenizer.encode(text))
        else:
            # Fallback: rough estimate
            return len(text.split()) * 1.3
    
    def extract_text_from_pdf(self, file_path: str) -> List[Dict]:
        """Extract text from PDF with page numbers"""
        try:
            reader = PdfReader(file_path)
            pages = []
            
            for page_num, page in enumerate(reader.pages, start=1):
                text = page.extract_text()
                if text.strip():
                    pages.append({
                        "page_number": page_num,
                        "text": text.strip()
                    })
            
            logger.info("Extracted %s pages from PDF", len(pages))
            return pages
        except Exception as e:
            logger.error("Error extracting PDF text: %s", e)
            return []
    
    def extract_text_from_docx(self, file_path: str) -> List[Dict]:
        """Extract text from DOCX"""
        try:
            doc = docx.Document(file_path)
            text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
            
            return [{
                "page_number": 1,
                "text": text
            }]
        except Exception as e:
            logger.error("Error extracting DOCX text: %s", e)
            return []
    
    def extract_text_from_txt(self, file_path: str) -> List[Dict]:
        """Extract text from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            
            return [{
                "page_number": 1,
                "text": text
            }]
        except Exception as e:
            logger.error("Error extracting TXT text: %s", e)
            return []
    
    def extract_text_from_excel(self, file_path: str) -> List[Dict]:
        """Extract text from Excel file"""
        try:
            # Try different Excel engines
            try:
                df = pd.read_excel(file_path, engine='openpyxl', sheet_name=None)
            except Exception as e:
                logger.debug("openpyxl engine failed, falling back to xlrd: %s", e)
                df = pd.read_excel(file_path, engine='xlrd', sheet_name=None)
            
            pages = []
            page_num = 1
            
            for sheet_name, sheet_data in df.items():
                text_content = [f"Sheet: {sheet_name}"]
                text_content.append(sheet_data.to_string())
                
                pages.append({
                    "page_number": page_num,
                    "text": '\n'.join(text_content)
                })
                page_num += 1
            
            logger.info("Extracted %s sheets from Excel", len(pages))
            return pages
        except Exception as e:
            logger.error("Error extracting Excel text: %s", e)
            return []
    
    def extract_text_from_bytes(self, content: bytes, file_type: str) -> str:
        """Extract text from file content bytes"""
        try:
            if file_type.lower() in ['docx', 'doc']:
                doc = docx.Document(BytesIO(content))
                full_text = []
                for para in doc.paragraphs:
                    full_text.append(para.text)
                return '\n'.join(full_text)
            
            elif file_type.lower() in ['xlsx', 'xls']:
                try:
                    df = pd.read_excel(BytesIO(content), engine='openpyxl', sheet_name=None)
                except Exception as e:
                    logger.debug("openpyxl engine failed, falling back to xlrd: %s", e)
                    df = pd.read_excel(BytesIO(content), engine='xlrd', sheet_name=None)
                
                text_content = []
                for sheet_name, sheet_data in df.items():
                    text_content.append(f"Sheet: {sheet_name}")
                    text_content.append(sheet_data.to_string())
                    text_content.append("\n")
                
                return '\n'.join(text_content)
            
            elif file_type.lower() == 'txt':
                return content.decode('utf-8')
            
            else:
                logger.error("Unsupported file type for bytes extraction: %s", file_type)
                return ""
                
        except Exception as e:
            logger.error("Error extracting text from bytes: %s", e)
            return ""
    
    def _validate_file_path(self, file_path: str, base_dir: Path = None) -> bool:
        """Validate file path — blocks directory traversal and ensures file is under base_dir."""
        try:
            path = Path(file_path).resolve()

            if not path.exists() or not path.is_file():
                logger.error("File does not exist or is not a file: %s", file_path)
                return False

            # Block traversal in original path before resolution
            if '..' in str(file_path).replace('\\', '/'):
                logger.error("Directory traversal attempt blocked: %s", file_path)
                return False

            # If base_dir given, confirm resolved path is underneath it
            if base_dir:
                base = Path(base_dir).resolve()
                try:
                    path.relative_to(base)
                except ValueError:
                    logger.error("Path escapes base directory: %s", file_path)
                    return False
                return True

            # Fallback: ensure file is inside an expected uploads directory
            expected_dirs = {'uploads', 'course_materials', 'materials'}
            if not any(part in expected_dirs for part in path.parts):
                logger.error("File not in expected upload directory: %s", file_path)
                return False

            return True
        except Exception as e:
            logger.error("Path validation error: %s", e)
            return False
    
    def extract_text(self, file_path: str, file_type: str) -> List[Dict]:
        """Extract text from file based on type"""
        # Validate file path for security
        if not self._validate_file_path(file_path):
            logger.error("File path validation failed: %s", file_path)
            return []
        
        file_type = file_type.lower()
        
        if file_type == 'pdf' or file_path.endswith('.pdf'):
            return self.extract_text_from_pdf(file_path)
        elif file_type in ['docx', 'doc'] or file_path.endswith(('.docx', '.doc')):
            return self.extract_text_from_docx(file_path)
        elif file_type in ['xlsx', 'xls'] or file_path.endswith(('.xlsx', '.xls')):
            return self.extract_text_from_excel(file_path)
        elif file_type == 'txt' or file_path.endswith('.txt'):
            return self.extract_text_from_txt(file_path)
        else:
            logger.error("Unsupported file type: %s", file_type)
            return []
    
    def chunk_text(self, text: str, chunk_size: int = None, overlap: int = None) -> List[str]:
        """Split text into overlapping chunks by tokens"""
        # Use configured values if not provided
        chunk_size = chunk_size or self.chunk_size
        overlap = overlap or self.chunk_overlap
        if not self.tokenizer:
            # Fallback: split by words
            words = text.split()
            chunks = []
            step = max(1, chunk_size - overlap)
            for i in range(0, len(words), step):
                chunk = " ".join(words[i:i + chunk_size])
                if chunk.strip():
                    chunks.append(chunk)
            return chunks

        # Split by tokens
        tokens = self.tokenizer.encode(text)
        chunks = []
        step = max(1, chunk_size - overlap)
        for i in range(0, len(tokens), step):
            chunk_tokens = tokens[i:i + chunk_size]
            chunk_text = self.tokenizer.decode(chunk_tokens)
            if chunk_text.strip():
                chunks.append(chunk_text)
        
        return chunks
    
    def process_document(self, material_id: int, file_path: str, file_type: str, 
                        course_code: str = None) -> Tuple[bool, int, str]:
        """
        Process a document: extract text, chunk, generate embeddings, and index
        Returns: (success, chunk_count, error_message)
        """
        try:
            logger.info("Processing document %s: %s", material_id, file_path)
            
            # Extract text
            pages = self.extract_text(file_path, file_type)
            if not pages:
                return False, 0, "Failed to extract text from document"
            
            # Process each page
            all_chunks = []
            chunk_index = 0
            
            for page_data in pages:
                page_num = page_data["page_number"]
                page_text = page_data["text"]
                
                # Chunk the page text
                chunks = self.chunk_text(page_text)
                
                for chunk_text in chunks:
                    token_count = int(self.count_tokens(chunk_text))
                    
                    all_chunks.append({
                        "chunk_index": chunk_index,
                        "text": chunk_text,
                        "page_number": page_num,
                        "token_count": token_count,
                        "material_id": material_id,
                        "source_file": os.path.basename(file_path),
                        "course_code": course_code or "general"
                    })
                    chunk_index += 1
            
            if not all_chunks:
                return False, 0, "No chunks generated from document"
            
            logger.info("Generated %s chunks", len(all_chunks))
            
            # Generate embeddings
            logger.info("[EMB] Generating embeddings...")
            texts = [chunk["text"] for chunk in all_chunks]
            embeddings = self.embedding_model.encode(texts, show_progress_bar=False)
            
            # Index in ChromaDB
            logger.info("Indexing in vector database...")
            ids = [f"mat{material_id}_chunk{chunk['chunk_index']}" for chunk in all_chunks]
            metadatas = [{
                "material_id": chunk["material_id"],
                "chunk_index": chunk["chunk_index"],
                "page_number": chunk["page_number"],
                "token_count": chunk["token_count"],
                "source_file": chunk["source_file"],
                "course_code": chunk["course_code"]
            } for chunk in all_chunks]
            
            self.collection.add(
                embeddings=embeddings.tolist(),
                documents=texts,
                metadatas=metadatas,
                ids=ids
            )
            
            logger.info("Successfully indexed %s chunks", len(all_chunks))
            return True, len(all_chunks), ""
            
        except Exception as e:
            error_msg = f"Error processing document: {str(e)}"
            logger.error(error_msg)
            return False, 0, error_msg
    
    def search(self, query: str, course_code: str = None, top_k: int = None) -> List[Dict]:
        """
        Search for relevant chunks using semantic similarity
        Returns list of results with text, metadata, and similarity scores
        """
        # Use configured value if not provided
        top_k = top_k or self.default_top_k
        try:
            logger.info("Searching for: %s...", query[:100])
            
            # Generate query embedding
            query_embedding = self.embedding_model.encode([query])[0]
            
            # Build filter
            where_filter = None
            if course_code:
                where_filter = {"course_code": course_code}
            
            # Search in ChromaDB
            results = self.collection.query(
                query_embeddings=[query_embedding.tolist()],
                n_results=top_k,
                where=where_filter
            )
            
            # Format results — filter by min_similarity (distance < threshold means more similar)
            min_similarity = self.min_similarity
            # Chroma default uses L2 distance. Convert to 0-1 similarity: 1/(1+distance)
            # This always stays between 0 and 1 regardless of distance value
            formatted_results = []
            if results and results['ids'] and len(results['ids'][0]) > 0:
                for i in range(len(results['ids'][0])):
                    distance = results['distances'][0][i] if 'distances' in results else None
                    similarity = round(1 / (1 + distance), 3) if distance is not None else 1.0
                    if similarity >= min_similarity:
                        formatted_results.append({
                            "id": results['ids'][0][i],
                            "text": results['documents'][0][i],
                            "metadata": results['metadatas'][0][i],
                            "distance": distance,
                            "similarity": similarity
                        })
                        logger.info("Chunk accepted: similarity=%s (distance=%.3f)", similarity, distance)
                    else:
                        logger.info("Chunk rejected: similarity=%s below threshold %s", similarity, min_similarity)
            
            logger.info("Found %s relevant chunks", len(formatted_results))
            return formatted_results
            
        except Exception as e:
            logger.error("Error searching: %s", e)
            return []
    
    def delete_material(self, material_id: int) -> bool:
        """Delete all chunks for a material from vector DB"""
        try:
            # Get all IDs for this material
            results = self.collection.get(
                where={"material_id": material_id}
            )
            
            if results and results['ids']:
                self.collection.delete(ids=results['ids'])
                logger.info("Deleted %s chunks for material %s", len(results['ids']), material_id)
                return True
            
            return False
        except Exception as e:
            logger.error("Error deleting material: %s", e)
            return False
    
    def build_rag_context(self, search_results: List[Dict], max_tokens: int = None) -> str:
        """Build context string from search results with source citations"""
        # Use configured value if not provided
        max_tokens = max_tokens or self.max_context_tokens
        context_parts = []
        total_tokens = 0
        
        for i, result in enumerate(search_results, 1):
            text = result['text']
            metadata = result['metadata']
            page = metadata.get('page_number', 'N/A')
            material_id = metadata.get('material_id', 'N/A')

            chunk_text = f"[Source {material_id}, Page {page}]\n{text}\n"
            chunk_tokens = int(self.count_tokens(chunk_text))

            # Always include at least one chunk even if it exceeds the token budget
            if total_tokens + chunk_tokens > max_tokens and context_parts:
                break

            context_parts.append(chunk_text)
            total_tokens += chunk_tokens
        
        return "\n---\n".join(context_parts)
    
    def get_stats(self) -> Dict:
        """Get RAG system statistics"""
        try:
            count = self.collection.count()
            return {
                "total_chunks": count,
                "collection_name": self.collection.name,
                "embedding_model": self.embedding_model.get_sentence_embedding_dimension()
            }
        except Exception as e:
            logger.error("Error getting stats: %s", e)
            return {}

# Global RAG service instance
rag_service = RAGService()
